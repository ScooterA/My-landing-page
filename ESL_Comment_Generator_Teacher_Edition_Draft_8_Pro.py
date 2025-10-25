#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ESL Comment Generator — V6 (Teacher Edition + Behaviour)
11 skills (1–10) + Specific Behaviour Issue dropdown (below sliders)
"""

import random
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from dataclasses import dataclass
from typing import Dict, List, Optional

def _choose(lines):
    """Return a random choice from a list or an empty string if the list is empty."""
    import random
    return random.choice(lines) if lines else ""

from typing import Dict, List, Optional

# ---------------- Core Logic ----------------

SKILLS = [
    "phonics","reading","writing","speaking","listening",
    "grammar","vocabulary","spelling","participation","homework","behaviour"
]

def band_for_score(score: int) -> str:
    if score <= 3:
        return "low"
    if score <= 7:
        return "mid"
    return "high"

def _pronouns(gender_choice: str):
    g = (gender_choice or "She").strip().lower()
    if g.startswith("h"):
        return ("He","His","him")
    return ("She","Her","her")

def apply_tone(sentence: str, tone: str) -> str:
    t = (tone or "neutral").lower()
    if t == "warm":
        return sentence.replace(".", ". 😊")
    if t == "direct":
        sentence = sentence.replace(" is ", " is clearly ").replace(" are ", " are clearly ")
        sentence = sentence.replace(" can ", " can consistently ")
    return sentence

CONNECTORS = [
    "Additionally, ", "Moreover, ", "In general, ", "As a result, ",
    "Overall, ", "It's great to see that ", "I'm pleased that ", "It’s wonderful that "
]

def opening_pool(name: str) -> List[str]:
    return [
        f"{name} has shown steady progress in learning.",
        f"{name} continues to develop strong English skills.",
        f"{name} demonstrates commitment and growing confidence in class activities.",
        f"{name} has made good progress across key areas of English.",
        f"{name} shows a positive attitude and consistent effort in class."
    ]

def closing_pool(name: str) -> List[str]:
    return [
        f"Keep up the great effort, {name}!",
        f"Fantastic progress — stay confident, {name}!",
        f"Excellent attitude and steady improvement, {name}!",
        f"Keep working hard and reaching for success, {name}!",
        f"Wonderful progress — keep it up, {name}!"
    ]

# --- Simple placeholder BANKS (replace with full version later) ---
BANKS = {skill: {
    "low": [f"Pronoun is learning {skill} and benefits from regular review."],
    "mid": [f"Pronoun is improving at {skill} and shows steady progress."],
    "high": [f"Pronoun demonstrates strong ability in {skill} and applies skills confidently."]
} for skill in SKILLS}

# --- Specific Behaviour Issues ---
BEHAVIOUR_ISSUES = {
    "Disruptive in class": "{pronoun_cap} sometimes becomes distracted or disrupts others. Gentle redirection helps {object_pronoun} focus better during lessons.",
    "Poor listening": "{pronoun_cap} has difficulty listening carefully at times and benefits from eye contact and short, clear reminders.",
    "Speaking in native language": "{pronoun_cap} occasionally uses {possessive_low} native language instead of English, so consistent reminders are helpful.",
    "Not completing homework": "{pronoun_cap} sometimes forgets homework tasks and needs a clear routine for checking and submitting work on time.",
    "Not raising hand / calling out": "{pronoun_cap} tends to call out answers without waiting, so reminders about turn-taking are helpful.",
    "Not paying attention": "{pronoun_cap} loses focus easily and benefits from sitting closer to the teacher and short breaks to refocus."
}

@dataclass
class V5Input:
    student_name: str
    scores: Dict[str, int]
    gender_choice: str
    teacher_name: str
    tone: str
    behaviour_issue: str
    seed: Optional[int] = None

def generate_comment(data: V5Input) -> str:
    if data.seed is not None:
        random.seed(data.seed)

    # Pronouns
    pronoun_cap, possessive_cap, obj = _pronouns(data.gender_choice)
    pronoun_low, possessive_low = pronoun_cap.lower(), possessive_cap.lower()

    # Safe-capitalize student display name (preserves rest of casing)
    raw_name = (data.student_name or "Student").strip()
    display_name = raw_name[0].upper() + raw_name[1:] if raw_name else "Student"

    # Opening + body
    opening = _choose(opening_pool(display_name))
    parts: List[str] = [opening]

    tone = data.tone.lower() if data.tone != "random" else random.choice(["neutral","warm","direct"])

    for skill, score in data.scores.items():
        band = band_for_score(score)
        bank = BANKS.get(skill, {})
        sent = _choose(bank.get(band, []))
        # Substitute pronouns
        sent = (sent
                .replace("Pronoun", pronoun_cap).replace("pronoun", pronoun_low)
                .replace("Possessive", possessive_cap).replace("possessive", possessive_low))
        sent = apply_tone(sent, tone)
        parts.append(sent)

    body = " ".join(parts).strip()

    # Behaviour note (its own paragraph, only if selected)
    behaviour_para = ""
    if data.behaviour_issue and data.behaviour_issue != "None":
        tmpl = BEHAVIOUR_ISSUES.get(data.behaviour_issue, "")
        if tmpl:
            behaviour_text = tmpl.format(
                pronoun_cap=pronoun_cap,
                pronoun_low=pronoun_low,
                possessive_low=possessive_low,
                object_pronoun=obj
            )
            behaviour_para = f"\n\nBehaviour Note: {behaviour_text}"

    # Closing + signature (each on clean lines; closing before signature)
    closing = _choose(closing_pool(display_name))
    signature = f"\n\n{data.teacher_name.strip()}" if data.teacher_name and data.teacher_name.strip() else ""

    # Final layout: body ⟶ (blank) ⟶ behaviour note (optional) ⟶ (blank) ⟶ closing ⟶ signature
    return f"{body}{behaviour_para}\n\n{closing}{signature}".strip()


# ---------------- GUI ----------------
class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("ESL Comment Generator — V6 (Teacher Edition + Behaviour)")
        self.geometry("1000x720")
        self._build_widgets()

    def _build_widgets(self):
        main = ttk.Frame(self, padding=10)
        main.pack(fill="both", expand=True)

        # Info section
        info = ttk.LabelFrame(main, text="Student & Settings", padding=10)
        info.pack(fill="x", pady=(0,10))
        ttk.Label(info, text="Student Name:").grid(row=0, column=0, sticky="w")
        self.var_name = tk.StringVar()
        ttk.Entry(info, textvariable=self.var_name, width=25).grid(row=0, column=1, padx=6)
        ttk.Label(info, text="Teacher:").grid(row=0, column=2, sticky="w")
        self.var_teacher = tk.StringVar()
        ttk.Entry(info, textvariable=self.var_teacher, width=25).grid(row=0, column=3, padx=6)
        ttk.Label(info, text="Gender:").grid(row=0, column=4, sticky="w")
        self.var_gender = tk.StringVar(value="She")
        ttk.Combobox(info, textvariable=self.var_gender, values=["She","He"], width=6, state="readonly").grid(row=0, column=5, padx=6)
        ttk.Label(info, text="Tone:").grid(row=0, column=6, sticky="w")
        self.var_tone = tk.StringVar(value="warm")
        ttk.Combobox(info, textvariable=self.var_tone, values=["neutral","warm","direct","random"], width=10, state="readonly").grid(row=0, column=7, padx=6)

        for i in range(8):
            info.columnconfigure(i, weight=1)

        # Sliders
        sliders = ttk.LabelFrame(main, text="Skills (1–10)", padding=10)
        sliders.pack(fill="x", pady=(0,10))
        self.slider_vars: Dict[str, tk.IntVar] = {}
        cols = 2

        for idx, skill in enumerate(SKILLS):
            row = idx // cols
            col = (idx % cols) * 3
            label = ttk.Label(sliders, text=skill.capitalize())
            label.grid(row=row*2, column=col, sticky="w", padx=(0,6))
            var = tk.IntVar(value=5)
            self.slider_vars[skill] = var
            s = ttk.Scale(sliders, from_=1, to=10, orient="horizontal", variable=var)
            s.grid(row=row*2, column=col+1, sticky="we", padx=(0,6))
            ttk.Label(sliders, textvariable=tk.StringVar(value=str(var.get()))).grid(row=row*2, column=col+2, sticky="w")
            scale = ttk.Frame(sliders)
            scale.grid(row=row*2+1, column=col+1, sticky="we")
            for n in range(1, 11):
                ttk.Label(scale, text=str(n), font=("Arial", 8)).pack(side="left", expand=True)

        for c in range(6):
            sliders.columnconfigure(c, weight=1)

        # Behaviour Issue dropdown BELOW sliders
        beh_frame = ttk.Frame(main, padding=(10,0))
        beh_frame.pack(fill="x", pady=(0,10))
        ttk.Label(beh_frame, text="Specific Behaviour Issue:").pack(side="left")
        self.var_behaviour_issue = tk.StringVar(value="None")
        ttk.Combobox(beh_frame, textvariable=self.var_behaviour_issue,
                     values=["None"] + list(BEHAVIOUR_ISSUES.keys()),
                     width=40, state="readonly").pack(side="left", padx=8)

        # Output
        out = ttk.LabelFrame(main, text="Generated Comment", padding=10)
        out.pack(fill="both", expand=True)
        self.txt = tk.Text(out, wrap="word", height=12)
        self.txt.pack(fill="both", expand=True)

        # Buttons
        btns = ttk.Frame(main)
        btns.pack(fill="x", pady=8)
        ttk.Button(btns, text="Generate", command=self.on_generate).pack(side="left")
        ttk.Button(btns, text="Copy", command=self.on_copy).pack(side="left", padx=6)
        ttk.Button(btns, text="Clear", command=self.on_clear).pack(side="left")
        ttk.Button(btns, text="Export DOCX", command=self.on_export_docx).pack(side="right")

        self.status = tk.StringVar(value="Ready.")
        ttk.Label(main, textvariable=self.status, anchor="w").pack(fill="x", pady=(4,0))

    # --- Functions ---
    def collect_scores(self) -> Dict[str, int]:
        return {k:int(v.get()) for k,v in self.slider_vars.items()}

    def build_input(self) -> V5Input:
        return V5Input(
            student_name=self.var_name.get().strip() or "Student",
            scores=self.collect_scores(),
            gender_choice=self.var_gender.get(),
            teacher_name=self.var_teacher.get(),
            tone=self.var_tone.get(),
            behaviour_issue=self.var_behaviour_issue.get(),
        )

    def on_generate(self):
        try:
            data = self.build_input()
            text = generate_comment(data)
            self.txt.delete("1.0", "end")
            self.txt.insert("1.0", text)
            self.status.set("Comment generated.")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate comment:\n{e}")
            self.status.set("Error.")

    def on_copy(self):
        text = self.txt.get("1.0", "end").strip()
        if not text: return
        self.clipboard_clear()
        self.clipboard_append(text)
        self.status.set("Copied to clipboard.")

    def on_clear(self):
        self.txt.delete("1.0", "end")
        self.status.set("Cleared.")

    def on_export_docx(self):
        from docx import Document
        text = self.txt.get("1.0","end").strip()
        if not text:
            messagebox.showinfo("Export","Generate a comment first."); return
        fname = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document","*.docx")])
        if not fname: return
        doc = Document()
        doc.add_heading("Student Comment", level=1)
        for para in text.split("\n\n"):
            doc.add_paragraph(para)
        doc.save(fname)
        messagebox.showinfo("Export", f"Saved: {fname}")

def main():
    app = App()
    app.mainloop()

if __name__ == "__main__":
    main()
